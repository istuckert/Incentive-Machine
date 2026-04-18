"""
Adds five blocks of H-2B research to Visa_Incentive_Machine_Master_Research.docx.

Insertion anchors (bottom-to-top to preserve indices):
  Add 5 — Ch16 after paragraph [498] (EPI comment on 2025 IFR)
  Add 4 — Ch15 after paragraph [438] (⚑ PENDING: Exact return-rate statistics)
  Add 3 — Ch06 after paragraph [180] (◆ CROSS-PROGRAM enforcement collapse)
  Add 2 — Ch04 after paragraph [129] (⚑ PENDING: EPI prevailing wage savings)
  Add 1 — Ch01 after paragraph [45]  (✓ SOURCE: DOL OFLC H-2B Selected Statistics)
"""

import copy, os
from docx import Document
from docx.oxml.ns import qn

DOCX_PATH = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "reference", "visa", "Visa_Incentive_Machine_Master_Research.docx"
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def find_para(doc, substring):
    for i, p in enumerate(doc.paragraphs):
        if substring in p.text:
            return i, p
    return None, None


def insert_after(doc, anchor_substring, elements, label=""):
    idx, _ = find_para(doc, anchor_substring)
    if idx is None:
        raise RuntimeError(f"Anchor not found: '{anchor_substring[:80]}'")
    body = doc.element.body
    anchor_elem = doc.paragraphs[idx]._element
    pos = list(body).index(anchor_elem)
    for offset, elem in enumerate(elements):
        body.insert(pos + 1 + offset, elem)
    print(f"  {label}: inserted {len(elements)} elements after paragraph [{idx}]")
    print(f"    anchor: {doc.paragraphs[idx].text[:80]}")


def harvest(tmp):
    return [
        copy.deepcopy(c)
        for c in tmp.element.body
        if c.tag != qn("w:sectPr")
    ]


def build(lines):
    """Build a list of XML elements from a list of strings."""
    tmp = Document()
    for line in lines:
        tmp.add_paragraph(line)
    return harvest(tmp)


# ---------------------------------------------------------------------------
# Content blocks
# ---------------------------------------------------------------------------

ADDITION_1 = [
    "H-2B TRUE PROGRAM SIZE (FY2024 — CONFIRMED)",
    "Statutory cap: 66,000 (set by Congress 1992, never raised by legislation)",
    "Supplemental cap authority: DHS authorized to add up to 64,716 additional visas per year via annual "
    "appropriations riders — both parties have authorized this every year since FY2016",
    "Total FY2024 cap: 130,716",
    "Total H-2B workers actually employed FY2024: 169,177 — more than 2.5x the original statutory cap",
    "Breakdown: 139,541 new visas + 4,580 extensions with same employer + 25,056 employer transfers",
    "✓ SOURCE: EPI analysis of USCIS H-2B Employer Data Hub + U.S. State Department Nonimmigrant Visa "
    "Statistics — epi.org/310379 + uscis.gov/tools/reports-and-studies/h-2b-employer-data-hub",
    "",
    "H-2B JOB DURATION — TEMPORARY IN NAME ONLY (FY2024)",
    "Average certified job duration: 232 days (7.6 months)",
    "86% of jobs certified for 6 months or more",
    "75% certified for 7 months or more",
    "53% certified for 8 months or more",
    "28% certified for 9 months or more — the legal maximum",
    "Program legally defined as temporary. Over half the certified positions run 8+ months.",
    "✓ SOURCE: DOL OFLC H-2B Performance Data FY2024 — dol.gov/agencies/eta/foreign-labor/performance",
    "",
    "H-2B OCCUPATION CONCENTRATION (FY2024)",
    "Top-10 occupations account for 98.8% of all H-2B approvals — program is highly concentrated in "
    "low-wage industries with documented wage theft histories",
    "#1 Building/grounds cleaning and maintenance (landscaping, maids, janitors): 93,349 approvals — 45.86%",
    "#2 Food preparation and serving: 28,746 approvals — 14.12%",
    "#3 Construction and extraction: 15,684 approvals — 7.70%",
    "#4 Production occupations: 15,476 approvals — 7.60%",
    "#5 Farming, fishing and forestry: 14,669 approvals — 7.21%",
    "✓ SOURCE: EPI analysis of USCIS H-2B Employer Data Hub FY2024 — epi.org/310379",
]

ADDITION_2 = [
    "H-2B CERTIFIED WAGES VS. NATIONAL STANDARDS — TOP OCCUPATIONS (FY2024)",
    "In all top-15 H-2B occupations, certified wages were lower than OEWS national average for the "
    "occupation. The DOL uses OEWS data to set H-2B wages — making this an apples-to-apples comparison "
    "showing the program legally suppresses national wage standards.",
    "Landscaping and groundskeeping (37% of all H-2B): H-2B avg $17.55/hr vs national avg $19.66/hr — 10.7% below",
    "Forest and conservation workers: H-2B avg $15.51/hr vs national avg $20.59/hr — 24.7% below (largest gap)",
    "Meat, poultry and fish cutters: H-2B avg $14.45/hr vs national avg $18.58/hr — 22.2% below",
    "Construction laborers: H-2B avg $20.85/hr vs national avg $24.64/hr — 15.4% below",
    "Cement masons: H-2B avg $23.66/hr vs national avg $28.54/hr — 17.1% below",
    "Fast food and counter workers: H-2B avg $12.68/hr vs national avg $15.07/hr — 15.8% below",
    "Packers and packagers: H-2B avg $14.85/hr vs national avg $17.59/hr — 15.6% below",
    "✓ SOURCE: EPI analysis of BLS 2024 OEWS data + DOL OFLC H-2B Performance Data FY2024 — epi.org/310379",
    "",
    "EMPLOYER PRIVATE WAGE SURVEY LOOPHOLE",
    "Employers can substitute their own private wage surveys to set H-2B minimum wages below even the "
    "local OEWS average. Seafood employers in particular use this loophole to pay workers significantly "
    "less than local or state averages require.",
    "✓ SOURCE: EPI September 2025 — epi.org/310379",
    "",
    "MEATPACKING H-2B WAGES VS. INDUSTRY (FY2024)",
    "Animal processing (meatpacking + poultry) employs roughly 560,000 people with combined payroll of "
    "nearly $30 billion. Foreign-born workers constitute 42.2% of the workforce.",
    "Average hourly wages: All workers $23.88 / U.S.-born $26.20 / Foreign-born $20.69 / "
    "New H-2B certifications $16.66",
    "U.S.-born meatpacking wages = 83.2% of typical U.S.-born economy-wide wages",
    "H-2B certified meatpacking wages = $16.66/hr — below even the already-suppressed foreign-born "
    "industry average",
    "✓ SOURCE: BLS Current Employment Statistics + Current Population Survey + DOL OFLC H-2B Performance "
    "Data FY2024 — epi.org/310379",
]

ADDITION_3 = [
    "H-2B INDUSTRY WAGE THEFT — AGGREGATE FY2000–2024 (inflation-adjusted to 2024 dollars)",
    "Source: DOL Wage and Hour Division, Industries with High Prevalence of H-2B Workers — enforcement "
    "covers ALL workers in these industries (U.S.-born, green card holders, H-2B workers, unauthorized "
    "workers), not H-2B workers exclusively. This means wage suppression from H-2B affects the entire "
    "industry workforce.",
    "Total compliance actions FY2000–2024: 247,180",
    "Total employees involved in violations: 2,021,557",
    "Total employees assessed back wages owed: 1,822,164",
    "Total back wages assessed: $2,238,327,967 (2024 dollars)",
    "Average back wages owed per employee: $1,228",
    "Total civil money penalties: $167,554,391",
    "Average stolen per year: $89.5 million annually across 25 fiscal years",
    "By industry: Construction $1,052,484,332 (avg $1,768/worker) / Food services $829,555,849 "
    "(avg $1,001/worker) / Janitorial $119,191,774 / Hotels and motels $105,638,105 / "
    "Landscaping $77,363,583 / Amusement $40,037,461 / Forestry $14,056,849",
    "✓ SOURCE: DOL Wage and Hour Division, Industries with High Prevalence of H-2B Workers, "
    "FY2000–2024 — dol.gov/agencies/whd",
    "",
    "WHD ENFORCEMENT CAPACITY COLLAPSE (Updated 2025)",
    "Only 611 WHD investigators as of May 2025 — policing a labor market of 170 million workers",
    "✓ SOURCE: Workplace Justice Lab, Rutgers University, May 2025 — cited in EPI September 2025 "
    "report epi.org/310379",
    "",
    "MEATPACKING EMPLOYER STRATEGY — DOCUMENTED (1983–2024)",
    "Historical pattern: As meatpacking unionization rates fell from 37.4% in 1983 to 13.5% in 2024, "
    "share of foreign-born workers rose to over 42%. This was an intentional employer strategy — "
    "aggressive recruitment of non-union immigrant labor to suppress worker bargaining power and wages.",
    "Current lobbying: Meatpacking employers are actively lobbying Congress to change the law to allow "
    "H-2B for year-round jobs. Currently prohibited because H-2B requires temporary work and meatpacking "
    "is not seasonal. The push is to convert precarious unauthorized workforce into an even more "
    "controlled tethered-visa workforce.",
    "This is the H-2 contractor layer mechanism extended into a new industry — visa dependency used to "
    "suppress wages and prevent organizing.",
    "✓ SOURCE: EPI analysis of Current Population Survey microdata 1983–2024 + David Calamuci, "
    "Return to the Jungle, New Labor Forum 2008 — cited in epi.org/310379",
]

ADDITION_4 = [
    "H-2B RESEARCH SUBSTANTIALLY COMPLETE (September 2025 EPI Report): Program size, occupation "
    "concentration, job duration, wage suppression by occupation, wage theft aggregate by industry, "
    "WHD enforcement capacity, meatpacking employer strategy — all now documented from EPI September "
    "2025 report (epi.org/310379) cross-referencing DOL OFLC, USCIS H-2B Employer Data Hub, "
    "BLS OEWS, and DOL WHD primary sources.",
]

ADDITION_5 = [
    "EPI September 2025 H-2B Report (Costa and Bivens) — epi.org/310379 / "
    "files.epi.org/uploads/the-h-2-b-visa-program-has-ballooned-without-being-fixed.pdf",
    "DOL WHD Industries with High Prevalence of H-2B Workers (FY2000–2024) — dol.gov/agencies/whd "
    "(wage theft enforcement data by industry)",
    "Workplace Justice Lab, Rutgers University, May 2025 — WHD investigator count 611 "
    "(cited in EPI September 2025)",
    "USCIS H-2B Employer Data Hub — uscis.gov/tools/reports-and-studies/h-2b-employer-data-hub",
]


# ---------------------------------------------------------------------------
# Main — apply in reverse order (bottom to top) to preserve indices
# ---------------------------------------------------------------------------

def main():
    print(f"Opening: {DOCX_PATH}")
    doc = Document(DOCX_PATH)
    print(f"  Paragraphs before: {len(doc.paragraphs)}")

    # Addition 5 — Ch16, after last EPI entry
    insert_after(
        doc,
        "EPI comment on 2025 IFR — epi.org/publication/epi-comment-on-dols-2025-interim-final-rule",
        build(ADDITION_5),
        label="Addition 5 (Ch16 sources)"
    )

    # Addition 4 — Ch15, after last Phase 3b H-2 pending item
    insert_after(
        doc,
        "⚑ PENDING: Exact return-rate statistics for H-2A workers across multiple seasons",
        build(ADDITION_4),
        label="Addition 4 (Ch15 Phase 3b)"
    )

    # Addition 3 — Ch06, after ◆ CROSS-PROGRAM enforcement paragraph
    insert_after(
        doc,
        "◆ CROSS-PROGRAM: Enforcement in both programs is collapsing relative to program scale",
        build(ADDITION_3),
        label="Addition 3 (Ch06 H-2 Enforcement)"
    )

    # Addition 2 — Ch04, after ⚑ PENDING: EPI prevailing wage savings paragraph
    insert_after(
        doc,
        "⚑ PENDING: EPI prevailing wage savings calculations per employer (Phase 4b)",
        build(ADDITION_2),
        label="Addition 2 (Ch04 H-2B wages)"
    )

    # Addition 1 — Ch01, after ✓ SOURCE: DOL OFLC H-2B Selected Statistics FY2025 Q4
    insert_after(
        doc,
        "✓ SOURCE: DOL OFLC H-2B Selected Statistics FY2025 Q4",
        build(ADDITION_1),
        label="Addition 1 (Ch01 H-2B program size)"
    )

    doc.save(DOCX_PATH)
    print(f"\n  Paragraphs after: {len(doc.paragraphs)}")
    print(f"  Saved: {DOCX_PATH}")


if __name__ == "__main__":
    main()
