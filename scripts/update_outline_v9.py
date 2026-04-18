"""
Updates Visa_Project_Outline_v9.docx with five sets of changes:
  1. Architecture update block after the title block
  2. Phase completion status table updates
  3. Section 04 status line replacement + verified Phase 4 data subsection
  4. Section 07 Phase 4 pending items replaced with completion block
  5. Decisions Log — April 2026 appended at end
"""

import copy
import os
from docx import Document
from docx.oxml.ns import qn

DOCX_PATH = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "reference", "visa", "Visa_Project_Outline_v9.docx"
)

# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------

def para_index(doc, substring):
    """Return index of first paragraph containing substring, or None."""
    for i, p in enumerate(doc.paragraphs):
        if substring in p.text:
            return i
    return None


def insert_elements_after_para(doc, para_idx, elements):
    """Insert a list of XML elements immediately after doc.paragraphs[para_idx]."""
    body = doc.element.body
    anchor = doc.paragraphs[para_idx]._element
    pos = list(body).index(anchor)
    for offset, elem in enumerate(elements):
        body.insert(pos + 1 + offset, elem)


def replace_para_elements_with(doc, start_idx, end_idx, elements):
    """
    Remove paragraphs[start_idx..end_idx] (inclusive) and insert elements
    in their place.
    """
    body = doc.element.body
    # Collect elements to remove
    to_remove = [doc.paragraphs[i]._element for i in range(start_idx, end_idx + 1)]
    anchor_pos = list(body).index(to_remove[0])
    for elem in to_remove:
        body.remove(elem)
    for offset, elem in enumerate(elements):
        body.insert(anchor_pos + offset, elem)


def harvest(tmp_doc):
    """Return deep-copied body children (minus sectPr) from a temp Document."""
    return [
        copy.deepcopy(child)
        for child in tmp_doc.element.body
        if child.tag != qn("w:sectPr")
    ]


# ---------------------------------------------------------------------------
# Content builders (each returns a list of XML elements)
# ---------------------------------------------------------------------------

def build_arch_update():
    tmp = Document()
    lines = [
        "ARCHITECTURE UPDATE — APRIL 2026",
        "Tool migrated from single HTML file to multi-page static site hosted on GitHub Pages.",
        "Live URL: https://istuckert.github.io/Incentive-Machine/",
        "Platform name changed to: The Manufactured Divide",
        "First module name: H-1B / H-2: The Visa Machine",
        'Tagline confirmed: "The argument is the product."',
        "Build environment: Plain HTML/CSS/JS — no framework, no build tools, no backend.",
    ]
    for line in lines:
        p = tmp.add_paragraph()
        run = p.add_run(line)
        if line == lines[0]:
            run.bold = True
    return harvest(tmp)


def build_phase4_data():
    tmp = Document()
    p = tmp.add_paragraph()
    p.add_run("VERIFIED PHASE 4 DATA (April 2026)").bold = True

    tmp.add_paragraph(
        "Big Three ownership — Top H-1B Sponsors (Q4 2025 13F):"
    ).runs[0].bold = True
    for line in [
        "Amazon: Vanguard 7.86% / BlackRock 6.83% / State Street 3.61%",
        "Meta: Vanguard 7.91% / BlackRock 6.78% / State Street 3.59%",
        "Microsoft: Vanguard 9.67% / BlackRock 8.11% / State Street 4.12%",
        "Google: Vanguard 7.74% / BlackRock 6.64% / State Street 3.44%",
        "TCS: Vanguard 0.77% / BlackRock 1.04% / State Street <0.5%",
    ]:
        tmp.add_paragraph(f"•  {line}")

    tmp.add_paragraph(
        "Federal lobbying 2025: Meta $26.29M / Amazon $18.87M / Google $16.54M / "
        "Microsoft $10.11M / TCS $1.04M"
    )
    tmp.add_paragraph(
        "PAC contributions 2023–2024: All four US sponsors split near 50/50 Dem/Rep "
        "— bipartisan access strategy"
    )
    tmp.add_paragraph(
        "Trump 2025 inaugural: Amazon, Meta, Microsoft, Google each donated $1M simultaneously"
    )
    tmp.add_paragraph(
        "Federal contracts: Amazon/AWS ~$798M / Microsoft ~$532M / Google low tens of millions "
        "/ Meta negligible / TCS minimal"
    )
    tmp.add_paragraph(
        "State subsidies: Amazon $11.6B+ / Google $2.319B / Microsoft $1.602B / Meta $60M+ / TCS none"
    )

    p = tmp.add_paragraph()
    p.add_run("TCS Consulting Scheme — Key Finding:").bold = True
    tmp.add_paragraph(
        "TCS acts as middleman — sponsors H-1B visas but deploys workers to end-client companies."
    )
    tmp.add_paragraph(
        "Documented end clients: Citigroup (largest, ~3,000 contractors 2020–2024), JPMorgan, "
        "Wells Fargo, Goldman Sachs, Capital One, Verizon, AT&T, Walmart, Barclays."
    )
    tmp.add_paragraph(
        "Big Three own all end-client companies at 5–10% stakes — same ownership loop extends to TCS clients."
    )

    p = tmp.add_paragraph()
    p.add_run("H-1B Approvals vs Layoffs — Same Companies Same Years:").bold = True
    for line in [
        "Amazon: 4,644 new H-1B FY2025 + ~30,000 US layoffs",
        "Meta: 1,555 new H-1B FY2025 (+112% vs FY2023) + ~21,000 layoffs",
        "Microsoft: 1,394 new H-1B FY2025 + ~15,000 layoffs same year",
        "Google: 1,050 new H-1B FY2025 + no major layoffs",
    ]:
        tmp.add_paragraph(f"•  {line}")

    return harvest(tmp)


def build_phase4_complete_block():
    tmp = Document()
    p = tmp.add_paragraph()
    p.add_run("PHASE 4 — COMPLETE (April 2026):").bold = True
    for line in [
        "✓ Per-company ownership chain confirmed — SEC 13F Q4 2025",
        "✓ Federal contract receipts — USASpending.gov",
        "✓ Lobbying expenditures — OpenSecrets 2025",
        "✓ PAC contributions — FEC 2023–2024 cycle",
        "✓ Layoff records vs H-1B hiring — WARN Act / CNBC / Reuters",
        "✓ TCS end-client documentation — Bloomberg 2025 investigation",
        "✓ State subsidies — Good Jobs First Subsidy Tracker",
        "✓ Federal tax paid — ITEP / Macrotrends / company 10-Ks",
        "✓ Trump 2025 inaugural donations — Guardian / CNBC",
        "✓ Job type breakdown per sponsor — myvisajobs.com / h1bscope.com",
        "REMAINING: H-2 PE intermediary chain (KKR/Blackstone) — still pending",
    ]:
        tmp.add_paragraph(f"•  {line}")
    return harvest(tmp)


def build_decisions_log():
    tmp = Document()
    p = tmp.add_paragraph()
    p.add_run("DECISIONS LOG — APRIL 2026").bold = True
    entries = [
        'Platform renamed from "The H-1B/H-2 Visa Incentive Machine" to "The Manufactured Divide"',
        'First module: "H-1B / H-2: The Visa Machine"',
        "Architecture migrated from single HTML to multi-page static site on GitHub Pages",
        "Home screen redesigned as D3 force-directed web with five module nodes",
        "Color scheme: cream/navy/blue (v8 palette)",
        "No center node on web — modules connect directly to each other",
        'Tagline confirmed: "The argument is the product."',
        "Jefferson 1820 quote added to hero with primary source link",
        "AI disclosure added — research by human, AI used for coding and organization",
        "Google Form feedback: https://forms.gle/fLaPAhALSbHqioBa6",
        "University of Kentucky College of Law removed from all attribution",
        "Credit: Created by Ian only",
        "Phase 4 research complete April 2026 — build pending",
        "Visa module rebuild deferred until Phase 6 research complete or near-complete",
        "H-2 contractor layer identified as new structural insight — two-tier H-2 chain: "
        "Sponsors (FLCs) → Operating Companies → Asset Managers",
        "TCS consulting scheme documented — Big Three own both outsourcers and end clients",
        "University feedback loop and corporate espionage sections remain architecturally orphaned "
        "— placement decision pending",
    ]
    for entry in entries:
        tmp.add_paragraph(f"•  {entry}")
    return harvest(tmp)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    print(f"Opening: {DOCX_PATH}")
    doc = Document(DOCX_PATH)

    # =========================================================
    # CHANGE 2 — Phase completion status table (TABLE 3)
    # Do this before any paragraph insertions so table index is stable.
    # =========================================================
    tbl = doc.tables[3]

    updates = {
        "Phase 1":   ("✓ Complete — migrated to multi-page static site", "v1–v3"),
        "Phase 2":   (
            "✓ Substantially Complete",
            "Content written, visuals built, sources verified. "
            "Vote flow diagram and foundation text added to visa/index.html. "
            "SEC 13F cross-ownership percentages confirmed."
        ),
        "Phase 3a":  (
            "🔶 Partial",
            "Sponsor table live. Scale, wage, rubber-stamp, and visa-dependency data now verified and slotted. "
            "Approval trend chart data now available (NFAP/USCIS 1992–2026)."
        ),
        "Phase 4":   (
            "✓ Research Complete — Build Pending",
            "All ownership chains, federal contracts, state subsidies, lobbying, PAC, inaugural donations, "
            "layoff data, job types, TCS scheme — documented in Master Research Chapter 09-A through 09-L."
        ),
    }

    for row in tbl.rows[1:]:  # skip header row
        phase_text = row.cells[0].text.strip()
        for key, (new_status, new_notes) in updates.items():
            if phase_text.startswith(key):
                row.cells[1].paragraphs[0].clear()
                row.cells[1].paragraphs[0].add_run(new_status)
                row.cells[2].paragraphs[0].clear()
                row.cells[2].paragraphs[0].add_run(new_notes)
                print(f"  Updated table row: {phase_text}")
                break

    # =========================================================
    # All remaining changes work bottom-to-top so indices stay valid.
    # Order: Change 5 → Change 4 → Change 3 → Change 1
    # =========================================================

    # =========================================================
    # CHANGE 5 — Decisions Log at end of document
    # =========================================================
    # Append to body (after last paragraph)
    last_para_idx = len(doc.paragraphs) - 1
    insert_elements_after_para(doc, last_para_idx, build_decisions_log())
    print(f"  Change 5: Decisions log appended after paragraph {last_para_idx}")

    # =========================================================
    # CHANGE 4 — Replace Phase 4 pending items in Section 07
    # =========================================================
    # Target paragraphs to remove/replace:
    #   "⚑ Pending: USASpending.gov — per-company federal contract data"
    #   "⚑ Pending: FEC.gov — PAC contribution data per company"
    #   "⚑ Pending: SEC 13F filings — cross-ownership exact percentages..."
    usa_idx    = para_index(doc, "USASpending.gov — per-company federal contract data")
    fec_idx    = para_index(doc, "FEC.gov — PAC contribution data per company")
    sec13f_idx = para_index(doc, "SEC 13F filings — cross-ownership exact percentages")

    if None in (usa_idx, fec_idx, sec13f_idx):
        raise RuntimeError(f"Change 4: Could not locate Phase 4 pending paragraphs. "
                           f"usa={usa_idx} fec={fec_idx} sec13f={sec13f_idx}")

    # They should be consecutive; replace the range
    start = min(usa_idx, fec_idx, sec13f_idx)
    end   = max(usa_idx, fec_idx, sec13f_idx)
    replace_para_elements_with(doc, start, end, build_phase4_complete_block())
    print(f"  Change 4: Replaced Phase 4 pending items (paragraphs {start}–{end})")

    # =========================================================
    # CHANGE 3 — Section 04 status line + verified data block
    # =========================================================
    # 3a: Replace the status sentence
    status_idx = para_index(doc, "Placeholder tables with FY2025 H-1B data populated")
    if status_idx is None:
        raise RuntimeError("Change 3: Could not find Section 04 status paragraph.")

    doc.paragraphs[status_idx].clear()
    doc.paragraphs[status_idx].add_run(
        "Status: Research complete as of April 2026. Build pending. "
        "All data documented in Master Research File Chapter 09-A through 09-L."
    )
    print(f"  Change 3a: Replaced status line at paragraph {status_idx}")

    # 3b: Insert verified Phase 4 data after the last Asset-Manager Tie-In bullet
    tie_in_last = para_index(doc, "This is the phase that closes the loop")
    if tie_in_last is None:
        raise RuntimeError("Change 3b: Could not find Asset-Manager Tie-In last bullet.")

    insert_elements_after_para(doc, tie_in_last, build_phase4_data())
    print(f"  Change 3b: Inserted verified Phase 4 data block after paragraph {tie_in_last}")

    # =========================================================
    # CHANGE 1 — Architecture update after title block
    # Insert after "One machine. Two visa programs. The same people get rich."
    # =========================================================
    tagline_idx = para_index(doc, "One machine. Two visa programs. The same people get rich.")
    if tagline_idx is None:
        raise RuntimeError("Change 1: Could not find tagline paragraph.")

    insert_elements_after_para(doc, tagline_idx, build_arch_update())
    print(f"  Change 1: Architecture update inserted after paragraph {tagline_idx}")

    # =========================================================
    # Save
    # =========================================================
    doc.save(DOCX_PATH)
    print(f"\nSaved successfully: {DOCX_PATH}")


if __name__ == "__main__":
    main()
