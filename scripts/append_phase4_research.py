"""
Appends Phase 4 research data (subsections 09-A through 09-L) into
Chapter 09 of the Master Research docx, and updates Chapter 15 with
a Phase 4 completion note.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import os

DOCX_PATH = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "reference", "visa", "Visa_Incentive_Machine_Master_Research.docx"
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def add_heading(doc, text, level=2):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    return p


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    # Data rows
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
    return table


def find_paragraph_index(doc, search_text):
    """Return index of first paragraph whose text contains search_text."""
    for i, p in enumerate(doc.paragraphs):
        if search_text in p.text:
            return i
    return None


def insert_paragraphs_after(doc, after_index, new_elements):
    """
    Insert a list of paragraph/table XML elements into doc.element.body
    immediately after the element at after_index in doc.paragraphs.
    Returns the index of the last inserted element.
    """
    body = doc.element.body
    anchor = doc.paragraphs[after_index]._element
    # Insert after anchor
    anchor_pos = list(body).index(anchor)
    for offset, elem in enumerate(new_elements):
        body.insert(anchor_pos + 1 + offset, elem)
    return after_index  # caller doesn't need the new index


# ---------------------------------------------------------------------------
# Build all new content as a temporary Document, then transplant its body
# elements into the real document at the right position.
# ---------------------------------------------------------------------------

def build_new_content():
    tmp = Document()

    # ---- 09-A ---------------------------------------------------------------
    add_heading(tmp, "09-A — BIG THREE OWNERSHIP: TOP H-1B SPONSORS (CONFIRMED Q4 2025)", level=2)
    add_para(tmp,
        "Ownership percentages confirmed from Q4 2025 SEC 13F filings, aggregated via "
        "Yahoo Finance, Wall Street Zen, and Investing.com. Cross-verified across multiple aggregators."
    )
    add_table(tmp,
        headers=["Company", "Vanguard", "BlackRock", "State Street", "Primary Source"],
        rows=[
            ["Amazon (AMZN)", "7.86%", "6.83%", "3.61%", "finance.yahoo.com/quote/AMZN/holders"],
            ["Meta (META)", "7.91%", "6.78%", "3.59%", "wallstreetzen.com/stocks/us/nasdaq/meta/ownership"],
            ["Microsoft (MSFT)", "9.67%", "8.11%", "4.12%", "finance.yahoo.com/quote/MSFT/holders"],
            ["Google/Alphabet (GOOG)", "7.74%", "6.64%", "3.44%", "finance.yahoo.com/quote/GOOG/holders"],
            ["TCS (TCS.NS)", "0.77%", "1.04%", "<0.5%", "investing.com/equities/tata-consultancy-services-ownership"],
        ]
    )
    add_para(tmp,
        "TCS is majority-owned by Tata Sons (~71.77% promoter stake). Big Three hold far smaller stakes "
        "than in U.S.-listed sponsors. The ownership loop is significantly tighter for Amazon/Meta/Microsoft/Google."
    )

    # ---- 09-B ---------------------------------------------------------------
    add_heading(tmp, "09-B — BIG THREE OWNERSHIP: TCS END-CLIENT COMPANIES", level=2)
    add_para(tmp,
        "TCS and similar outsourcing firms deploy H-1B workers to third-party end-clients under service "
        "contracts. The Big Three own both the outsourcers AND their end clients. Largest documented TCS "
        "end-client: Citigroup (~3,000 new H-1B contractors placed 2020–2024). "
        "Source: bloomberg.com/graphics/2025-h1b-visa-middlemen-cheap-labor-for-us-banks"
    )
    add_table(tmp,
        headers=["End Client", "Vanguard", "BlackRock", "State Street", "Source"],
        rows=[
            ["Citigroup (C)", "8.9%", "7.2%", "4.3%", "finance.yahoo.com/quote/C/holders"],
            ["JPMorgan Chase (JPM)", "9.91%", "7.89%", "4.67%", "finance.yahoo.com/quote/JPM/holders"],
            ["Wells Fargo (WFC)", "9.61%", "8.50%", "4.44%", "finance.yahoo.com/quote/WFC/holders"],
            ["Goldman Sachs (GS)", "9.84%", "7.88%", "6.63%", "finance.yahoo.com/quote/GS/holders"],
            ["Capital One (COF)", "9.15%", "8.20%", "4.48%", "finance.yahoo.com/quote/COF/holders"],
            ["Verizon (VZ)", "9.08%", "8.83%", "5.33%", "finance.yahoo.com/quote/VZ/holders"],
            ["AT&T (T)", "9.51%", "8.16%", "4.76%", "finance.yahoo.com/quote/T/holders"],
            ["Walmart (WMT)", "5.52%", "4.38%", "2.32%", "finance.yahoo.com/quote/WMT/holders"],
            ["Barclays (BCS ADR)", "7.8%", "6.5%", "3.5%", "finance.yahoo.com/quote/BCS/holders"],
            ["USAA", "N/A", "N/A", "N/A", "Member-owned mutual — not publicly traded"],
        ]
    )
    add_para(tmp,
        "The Big Three own both the outsourcing firm (TCS) and every major end-client company receiving "
        "the workers. Labor cost savings flow to margins at every node — and the Big Three capture those "
        "margin improvements as index fund returns."
    )

    # ---- 09-C ---------------------------------------------------------------
    add_heading(tmp, "09-C — FEDERAL CONTRACTS", level=2)
    add_table(tmp,
        headers=["Company", "Federal Contract Amount", "Notes", "Source"],
        rows=[
            ["Amazon/AWS", "~$798M cumulative FY2022–2025", "Primarily cloud services", "usaspending.gov"],
            ["Microsoft", "~$472–532M", "Cloud/Azure/IT services", "usaspending.gov/recipient/dd77b7c3-663e-cb91-229f-5766a50e9b7f-P/all"],
            ["Google/Alphabet", "Low tens of millions", "Discount strategy; individual DoD awards $15.9M $13.2M $3.7M", "usaspending.gov"],
            ["Meta", "Negligible (<$50K tracked)", "No significant federal contracting footprint", "usaspending.gov"],
            ["TCS", "Scattered small contracts only", "Low millions cumulative; no major aggregate", "usaspending.gov"],
        ]
    )

    # ---- 09-D ---------------------------------------------------------------
    add_heading(tmp, "09-D — STATE SUBSIDIES (GOOD JOBS FIRST)", level=2)
    add_table(tmp,
        headers=["Company", "Cumulative Subsidies", "Primary Mechanism", "Source"],
        rows=[
            ["Amazon", "$11.6B+ since ~2000", "Warehouses, data centers, HQ — property tax abatements, sales tax exemptions, job credits", "goodjobsfirst.org/amazon-tracker"],
            ["Google/Alphabet", "$2.319B (131 awards)", "Data centers — Oklahoma $917M, Oregon $533M, NC $256M", "subsidytracker.goodjobsfirst.org/parent/alphabet-inc"],
            ["Microsoft", "$1.602B (100 awards) + $840M loans", "Data centers — property tax abatements, sales/use tax exemptions", "subsidytracker.goodjobsfirst.org/parent/microsoft"],
            ["Meta", "$60M+ tracked", "Data centers — WA $31M, OR $28M, TX and LA 80% property tax breaks", "subsidytracker.goodjobsfirst.org/parent/meta-platforms-inc"],
            ["TCS", "None identified", "No material US state/local subsidies tracked", "subsidytracker.goodjobsfirst.org/parent/tata-group"],
        ]
    )

    # ---- 09-E ---------------------------------------------------------------
    add_heading(tmp, "09-E — FEDERAL TAX PAID 2025", level=2)
    add_para(tmp,
        "Big Tech group collectively paid ~4.9% effective federal rate on $315B US profits in 2025, "
        "avoiding an estimated $51B+ in federal taxes. "
        "Source: itep.org/trump-meta-tesla-alphabet-amazon-obbba-taxes"
    )
    add_table(tmp,
        headers=["Company", "Federal Tax Paid", "Effective Rate", "Source"],
        rows=[
            ["Microsoft", "$21.795B provision", "~18%", "macrotrends.net/stocks/charts/MSFT/microsoft/total-provision-income-taxes"],
            ["Google/Alphabet", "$26.656B provision", "Low teens", "macrotrends.net/stocks/charts/GOOG/alphabet/total-provision-income-taxes"],
            ["Amazon", "$2.75–2.8B cash paid", "Sharp drop from prior years", "politico.com/news/2026/02/06/amazon-emerges-a-big-winner-from-gop-tax-cuts-00768985"],
            ["Meta", "~$2.8B cash paid", "~3.5–3.6% on $79B US income — avoided ~$13.7B", "itep.org/meta-tax-breaks-trump-mark-zuckerberg"],
            ["TCS", "Not broken out", "~24.5% global effective rate", "finbox.com/NSEI:TCS/explorer/effect_tax_rate"],
        ]
    )

    # ---- 09-F ---------------------------------------------------------------
    add_heading(tmp, "09-F — LOBBYING SPEND 2025", level=2)
    add_table(tmp,
        headers=["Company", "2025 Federal Lobbying", "Source"],
        rows=[
            ["Meta", "$26,290,000", "opensecrets.org/federal-lobbying/clients/summary?id=D000033563"],
            ["Amazon", "$18,865,000", "opensecrets.org/federal-lobbying/clients/summary?cycle=2025&id=D000023883"],
            ["Google/Alphabet", "$16,540,000", "opensecrets.org/federal-lobbying/clients/summary?id=D000067823"],
            ["Microsoft", "$10,105,000", "opensecrets.org/federal-lobbying/clients/summary?id=D000000115"],
            ["TCS/Tata Sons", "$1,040,000 (2024)", "opensecrets.org/orgs/tata-sons-ltd/summary?id=D000067061"],
        ]
    )
    add_para(tmp, "Total top 4 US tech sponsors: $71.8M in federal lobbying in 2025 alone.")

    # ---- 09-G ---------------------------------------------------------------
    add_heading(tmp, "09-G — PAC CONTRIBUTIONS 2023–2024 CYCLE", level=2)
    add_table(tmp,
        headers=["Company", "Total PAC", "Dem %", "Rep %", "Source"],
        rows=[
            ["Amazon", "$799,000", "50.56%", "49.44%", "opensecrets.org/political-action-committees-pacs/amazon-com/C00360354/summary/2024"],
            ["Microsoft", "$800,500", "45.16%", "53.40%", "opensecrets.org/political-action-committees-pacs/microsoft-corp/C00227546/summary/2024"],
            ["Google", "$737,066", "48.51%", "50.47%", "opensecrets.org/political-action-committees-pacs/google-inc/C00428623/summary/2024"],
            ["Meta", "$197,300", "40%", "57%", "opensecrets.org/political-action-committees-pacs/meta/C00502906/summary/2024"],
            ["TCS", "None", "N/A", "N/A", "No corporate PAC registered"],
        ]
    )
    add_para(tmp,
        "All four US tech sponsors split PAC contributions near 50/50. Bipartisan access strategy — "
        "no dominant party alignment. Reinforces Phase 6 political flip argument: same companies fund "
        "both sides simultaneously."
    )

    # ---- 09-H ---------------------------------------------------------------
    add_heading(tmp, "09-H — 2025 TRUMP INAUGURAL DONATIONS", level=2)
    add_para(tmp,
        "All four major H-1B sponsors donated $1M each to the Trump 2025 inaugural committee simultaneously "
        "— buying access to the administration overseeing H-1B policy during active regulatory uncertainty "
        "($100K fee, new wage rules, lottery changes all pending). Inaugural committees have no contribution "
        "limits — corporations can donate directly unlike campaign finance."
    )
    add_table(tmp,
        headers=["Company", "Amount", "Source"],
        rows=[
            ["Amazon", "$1,000,000 cash + $1M in-kind Prime Video streaming", "theguardian.com/technology/2024/dec/13/amazon-donation-trump-inauguration"],
            ["Meta", "$1,000,000 cash", "cnbc.com/2025/04/23/trump-inauguration-donors-include-meta-amazon-target-delta-ford.html"],
            ["Microsoft", "$1,000,000 cash", "theguardian.com/technology/2025/jan/09/google-microsoft-donate-trump-inaugural-fund"],
            ["Google", "$1,000,000 cash", "cnbc.com/2025/01/09/google-donates-1-million-to-trumps-inauguration-fund.html"],
            ["TCS", "None identified", "washingtonpost.com/politics/interactive/2025/trump-inauguration-donors-list"],
        ]
    )

    # ---- 09-I ---------------------------------------------------------------
    add_heading(tmp, "09-I — H-1B APPROVALS VS. LAYOFFS", level=2)
    add_table(tmp,
        headers=["Company", "H-1B New Approvals FY2025", "Layoffs Same Period", "Pattern"],
        rows=[
            ["Amazon", "4,644", "~30,000 corporate cuts Oct 2025 + Jan 2026", "H-1B up 20% during largest-ever workforce reduction"],
            ["Meta", "1,555 (+112% vs FY2023)", "~21,000 (2022–2023) + 8,000 planned May 2026", "H-1B doubled during year of efficiency cuts"],
            ["Microsoft", "1,394", "~15,000 (6,000 May + 9,000 July 2025)", "~1,400 new H-1B alongside 15,000 layoffs same year"],
            ["Google", "1,050", "Minimal — no major rounds 2023–2025", "Consistent H-1B with no equivalent layoff pattern"],
            ["TCS", "846 (down 42% from FY2024)", "No major US layoffs tracked", "Shift away from outsourcing model"],
        ]
    )
    add_para(tmp,
        "Sources: cnbc.com/2026/01/28/amazon-layoffs-anti-bureaucracy-ai | "
        "reuters.com/world/meta-targets-may-20-first-wave-layoffs | "
        "cnbc.com/2025/07/02/microsoft-laying-off-about-9000-employees"
    )

    # ---- 09-J ---------------------------------------------------------------
    add_heading(tmp, "09-J — JOB TYPES FILLED BY H-1B", level=2)
    add_para(tmp,
        "Source: DOL OFLC LCA data FY2020–FY2026 via myvisajobs.com, h1bscope.com, h1bexposed.tech"
    )
    add_para(tmp,
        "Amazon top roles: Software Development Engineer II (19,723), SDE I (13,183), SDE III (4,812), "
        "BI Engineer II (3,429). Education: 25.2% Bachelor, 71.2% Master, 3.6% Doctorate."
    )
    add_para(tmp,
        "Meta top roles: Software Engineer (15,629), Research Scientist (2,716), Data Scientist (2,188), "
        "Data Engineer (1,749). Education: 22.2% Bachelor, 71.4% Master, 6.4% Doctorate."
    )
    add_para(tmp,
        "Microsoft top roles: Software Engineering (9,542), Software Engineer (7,218), "
        "Product Management (806), Data Science (656). Education: 35.2% Bachelor, 61.3% Master, 3.5% Doctorate."
    )
    add_para(tmp,
        "Google top roles: Software Engineer (17,126), Program Manager (597), Product Manager (584), "
        "Research Scientist (437). Education: 25.2% Bachelor, 66.1% Master, 8.7% Doctorate."
    )
    add_para(tmp,
        "TCS top roles: Architect (10,291), Developer (4,030), Analyst (2,413), Technical Lead (1,368), "
        "System Administrator (812). Education: 85.2% Bachelor, 14.8% Master, 0.0% Doctorate."
    )
    add_para(tmp,
        "Important distinction: Direct employers (Amazon/Meta/Microsoft/Google) predominantly hire "
        "Master's holders in genuine specialty roles. Wage suppression operates through LCA level "
        "manipulation and reduced organizing capacity — not simple labor replacement. Outsourcing firms "
        "(TCS/Cognizant/Infosys) are Bachelor's-heavy with standardized implementation roles deployed "
        "to client sites. Cleaner displacement argument. More documented wage manipulation. "
        "Different mechanism, same program."
    )

    # ---- 09-K ---------------------------------------------------------------
    add_heading(tmp, "09-K — THE TCS CONSULTING SCHEME", level=2)
    add_para(tmp,
        "TCS and similar outsourcing firms (Infosys, Cognizant, Wipro, HCL, Capgemini) sponsor H-1B visas "
        "but act as middlemen — deploying workers to third-party end-client companies under service contracts "
        "rather than employing them in-house. Workers perform routine IT implementation and support at client "
        "facilities, achieving 30–40% labor cost savings vs direct US hires."
    )
    add_para(tmp,
        "Documented end clients per Bloomberg 2025 investigation: Citigroup (largest — ~3,000 new H-1B "
        "contractors placed 2020–2024), Capital One, Verizon, AT&T, Walmart, USAA, JPMorgan Chase, "
        "Wells Fargo, Goldman Sachs, Barclays, and broadly Fortune 500 corporations in financial services, "
        "telecom, insurance, retail, healthcare, and manufacturing."
    )
    add_para(tmp,
        "Sources: bloomberg.com/graphics/2025-h1b-visa-middlemen-cheap-labor-for-us-banks | "
        "redbus2us.com/top-20-clients-for-major-outsourcing-mncs-onsite-h1bs-dol-stats | "
        "epi.org/blog/new-data-infosys-tata-abuse-h-1b-program"
    )

    # ---- 09-L ---------------------------------------------------------------
    add_heading(tmp, "09-L — NARRATIVE INFRASTRUCTURE (PRO-EXPANSION THINK TANKS)", level=2)
    add_para(tmp,
        "The following organizations regularly analyze USCIS/DOL data and publish reports framing H-1B as "
        "economically beneficial. All are structured as 501(c)(3)s. Funding flows from tech-aligned foundations "
        "creating the same incentive loop: funders benefit from corporate visa model; research protects it."
    )
    add_para(tmp,
        "NFAP (National Foundation for American Policy) — Funded partly by FWD.us Education Fund "
        "(Zuckerberg immigration advocacy), Carnegie Corporation, ImpactAssets. "
        "ProPublica EIN 20-0094633: projects.propublica.org/nonprofits/organizations/200094633"
    )
    add_para(tmp,
        "Migration Policy Institute — Carnegie Corporation $2M grant 2026 plus Annie E. Casey, "
        "Arnold Ventures. Funders page: migrationpolicy.org/about/funders"
    )
    add_para(tmp,
        "Cato Institute — Koch network historically prominent; diversified free-market foundations. "
        "Primary H-1B voice: David Bier. cato.org/people/david-bier"
    )
    add_para(tmp,
        "American Immigration Council — 501(c)(3) focused on immigration policy education. "
        "Pro-expansion fact sheets on H-1B and H-2 programs. americanimmigrationcouncil.org"
    )
    add_para(tmp,
        "Economic Innovation Group (EIG) — 2026 report: H-1B households generate $30K+ net annual "
        "fiscal surplus. Tech and innovation funders. eig.org/fiscal-impacts-h1bs"
    )
    add_para(tmp,
        "For H-2 programs: American Farm Bureau Federation, American Immigration Council, Cato Institute, "
        "National Council of Agricultural Employers, H-2B Workforce Coalition (AHLA/Edgeworth Economics "
        "commissioned studies claiming H-2B creates 3–5 local jobs per visa worker). Pattern identical "
        "to H-1B: same USCIS/DOL raw data interpreted to emphasize benefits."
    )

    return tmp


def build_ch15_update():
    tmp = Document()
    add_para(tmp,
        "PHASE 4 STATUS UPDATE (April 2026): The following Phase 4 items are now COMPLETE and documented "
        "in Chapter 09 subsections A through L: Big Three ownership percentages for top H-1B sponsors (09-A); "
        "Big Three ownership of TCS end-client companies (09-B); Federal contract receipts per sponsor (09-C); "
        "State subsidies per sponsor (09-D); Federal tax paid per sponsor (09-E); Lobbying expenditures per "
        "sponsor (09-F); PAC contributions per sponsor (09-G); Trump inaugural donations (09-H); H-1B approvals "
        "vs layoff records (09-I); Job types filled per sponsor (09-J); TCS consulting/contractor scheme "
        "documentation (09-K); Narrative infrastructure documentation (09-L). "
        "REMAINING PHASE 4 GAPS: H-2 sponsor PE intermediary chain (KKR/Blackstone confirmation still pending). "
        "PHASE 6 research (vote records, donor tracking, legislator quotes, party platform language, "
        "Pew polling) remains outstanding.",
        bold=False
    )
    return tmp


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    print(f"Opening: {DOCX_PATH}")
    doc = Document(DOCX_PATH)

    # ---- Locate Chapter 09 boundary -----------------------------------------
    # Find the last paragraph inside Ch09 (just before Ch10 header)
    ch10_idx = find_paragraph_index(doc, "10 · THE UNIVERSITY FEEDBACK LOOP")
    if ch10_idx is None:
        raise RuntimeError("Could not find Chapter 10 heading — check document structure.")

    # The insertion point is the paragraph just before Ch10 header
    insert_after_idx = ch10_idx - 1
    print(f"  Ch10 found at paragraph index {ch10_idx}; inserting new Ch09 content after index {insert_after_idx}")

    # Build new Ch09 content in a temp doc, harvest its body elements
    tmp_09 = build_new_content()
    new_elements_09 = [
        copy.deepcopy(child)
        for child in tmp_09.element.body
        if child.tag != qn("w:sectPr")  # skip section properties
    ]

    # Insert into real doc before Ch10
    body = doc.element.body
    anchor = doc.paragraphs[insert_after_idx]._element
    anchor_pos = list(body).index(anchor)
    for offset, elem in enumerate(new_elements_09):
        body.insert(anchor_pos + 1 + offset, elem)
    print(f"  Inserted {len(new_elements_09)} elements for subsections 09-A through 09-L")

    # ---- Locate Chapter 15 boundary -----------------------------------------
    # Re-index paragraphs since we just inserted elements
    ch15_idx = find_paragraph_index(doc, "15 · PENDING GAPS")
    ch16_idx = find_paragraph_index(doc, "16 · VERIFIED SOURCES")
    if ch15_idx is None:
        raise RuntimeError("Could not find Chapter 15 heading.")
    if ch16_idx is None:
        raise RuntimeError("Could not find Chapter 16 heading.")

    insert_after_ch15 = ch16_idx - 1
    print(f"  Ch15 at {ch15_idx}, Ch16 at {ch16_idx}; appending Phase 4 status after index {insert_after_ch15}")

    tmp_15 = build_ch15_update()
    new_elements_15 = [
        copy.deepcopy(child)
        for child in tmp_15.element.body
        if child.tag != qn("w:sectPr")
    ]

    anchor2 = doc.paragraphs[insert_after_ch15]._element
    anchor_pos2 = list(doc.element.body).index(anchor2)
    for offset, elem in enumerate(new_elements_15):
        doc.element.body.insert(anchor_pos2 + 1 + offset, elem)
    print(f"  Inserted {len(new_elements_15)} element(s) for Ch15 Phase 4 status update")

    # ---- Save ---------------------------------------------------------------
    doc.save(DOCX_PATH)
    print(f"\nSaved successfully: {DOCX_PATH}")


if __name__ == "__main__":
    main()
