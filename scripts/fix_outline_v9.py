"""
Two targeted fixes to Visa_Project_Outline_v9.docx:

Fix 1 — Build Rules: Update "Single HTML file" bullet to reflect GitHub Pages architecture.

Fix 2 — Section 04 stale pending items: The "Per-Company Research Needed (Phase 4)" block
(paragraphs 217–228) is now superseded. Replace the heading with a completion note and
delete all the stale pending bullets beneath it.

Note: "Phase 4 — Asset Managers" and "SEC 13F filings — confirm exact..." paragraphs
described in the brief do not exist in this outline file (they existed only in the
Master Research file). The script reports this and proceeds with what is present.
"""

import os
from docx import Document

DOCX_PATH = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "reference", "visa", "Visa_Project_Outline_v9.docx"
)

COMPLETION_TEXT = (
    "Phase 4 — Corporate Layer: COMPLETE April 2026 — see Chapter 09-A through 09-L "
    "in Master Research File. One remaining gap: H-2 PE intermediary chain "
    "(KKR/Blackstone) — confirmation pending."
)


def find_para(doc, substring):
    """Return (index, paragraph) for first paragraph containing substring, or (None, None)."""
    for i, p in enumerate(doc.paragraphs):
        if substring in p.text:
            return i, p
    return None, None


def delete_para(doc, para):
    """Remove a paragraph element from the document body."""
    p_elem = para._element
    p_elem.getparent().remove(p_elem)


def main():
    print(f"Opening: {DOCX_PATH}")
    doc = Document(DOCX_PATH)

    changed = []

    # =========================================================================
    # FIX 1 — Build Rules: replace "Single HTML file" bullet
    # =========================================================================
    fix1_idx, fix1_para = find_para(doc, "Single HTML file — no install, no login, no internet required")
    if fix1_para is None:
        print("  FIX 1: WARNING — target paragraph not found.")
    else:
        fix1_para.clear()
        fix1_para.add_run(
            "•  Hosted multi-page static site on GitHub Pages — no install, no login required. "
            "All pages accessible via URL. Works in any browser."
        )
        print(f"  FIX 1: Replaced Build Rules 'Single HTML file' bullet at paragraph {fix1_idx}")
        changed.append("Fix 1 — Build Rules bullet updated")

    # =========================================================================
    # FIX 2 — Section 04: remove stale Phase 4 pending block
    #
    # Target the "Per-Company Research Needed (Phase 4)" heading and all
    # paragraphs that follow it until the next section (SECTION 05).
    # Then replace that heading with the completion note.
    #
    # Stale paragraphs to delete (search strings):
    STALE_SNIPPETS = [
        "4a: Visa count and year-over-year trend",
        "4b: Estimated labor cost savings",
        "4c: Ultimate ownership chain",
        "4d: Federal contract receipts",
        "4e: Lobbying expenditures — OpenSecrets",
        "4f/4g: PAC contributions + layoff juxtaposition",
        "⚑ Pending: USASpending.gov per-company contract data",
        "⚑ Pending: OpenSecrets lobbying per company",
        "⚑ Pending: FEC PAC contribution data",
        "⚑ Pending: EPI wage savings calculations per employer",
        "⚑ Pending: Documented layoff records by company and year",
    ]

    # Paragraphs to delete must be collected BEFORE any deletion
    # (deleting shifts indices, but we hold references to element objects).
    paras_to_delete = []
    for snippet in STALE_SNIPPETS:
        _, p = find_para(doc, snippet)
        if p is not None:
            paras_to_delete.append((snippet, p))
        else:
            print(f"  FIX 2: NOTE — snippet not found (already removed?): '{snippet[:60]}'")

    # Delete stale paragraphs
    for snippet, p in paras_to_delete:
        delete_para(doc, p)
        print(f"  FIX 2: Deleted paragraph containing: '{snippet[:70]}'")

    # Replace the heading paragraph
    _, heading_para = find_para(doc, "Per-Company Research Needed (Phase 4)")
    if heading_para is None:
        print("  FIX 2: WARNING — 'Per-Company Research Needed (Phase 4)' heading not found.")
    else:
        heading_para.clear()
        heading_para.add_run(COMPLETION_TEXT)
        print(f"  FIX 2: Replaced heading with Phase 4 completion note")
        changed.append("Fix 2 — Section 04 stale pending block replaced with completion note")

    # Check for "Phase 4 — Asset Managers" and SEC 13F confirm paragraph
    # (these exist in the Master Research file but NOT in this outline — report accordingly)
    for check in [
        "Phase 4 — Asset Managers",
        "SEC 13F filings — confirm exact Vanguard/BlackRock/State Street cross-ownership",
    ]:
        idx, p = find_para(doc, check)
        if p is not None:
            delete_para(doc, p)
            print(f"  FIX 2: Deleted paragraph containing: '{check}'")
            changed.append(f"Fix 2 — Deleted: '{check}'")
        else:
            print(f"  FIX 2: NOTE — '{check}' not present in this file (outline only, not Master Research) — skipping.")

    # =========================================================================
    # Save
    # =========================================================================
    doc.save(DOCX_PATH)
    print(f"\nSaved successfully: {DOCX_PATH}")
    print(f"\nChanges made ({len(changed)}):")
    for c in changed:
        print(f"  • {c}")


if __name__ == "__main__":
    main()
