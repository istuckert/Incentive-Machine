"""
Adds H-2B research data to four sections of Visa_Project_Outline_v9.docx.
Applied bottom-to-top to preserve paragraph indices.

Anchors:
  Change 4 — [338] "Centro de los Derechos del Migrante — cdmigrante..."
  Change 3 — [180] "2024 Farmworker Protection Rule suspended June 2025"
  Change 2 — [167] "EPI: 2026 IFR wage cuts transfer $1.7–$2.1 billion"
  Change 1 — [148] "H-2B FY2026 Q1: ~39k certified"
"""

import copy, os
from docx import Document
from docx.oxml.ns import qn

DOCX_PATH = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "reference", "visa", "Visa_Project_Outline_v9.docx"
)


def find_para(doc, substring):
    for i, p in enumerate(doc.paragraphs):
        if substring in p.text:
            return i, p
    return None, None


def insert_after(doc, anchor, lines, label):
    idx, _ = find_para(doc, anchor)
    if idx is None:
        raise RuntimeError(f"Anchor not found: '{anchor[:80]}'")
    tmp = Document()
    for line in lines:
        tmp.add_paragraph(line)
    elements = [
        copy.deepcopy(c) for c in tmp.element.body
        if c.tag != qn("w:sectPr")
    ]
    body = doc.element.body
    anchor_elem = doc.paragraphs[idx]._element
    pos = list(body).index(anchor_elem)
    for offset, elem in enumerate(elements):
        body.insert(pos + 1 + offset, elem)
    print(f"  {label}: {len(elements)} elements after [{idx}]")
    print(f"    anchor: {doc.paragraphs[idx].text[:90]}")


CHANGE_1 = [
    "•  H-2B true program size FY2024: 169,177 workers actually employed — more than 2.5x the original "
    "statutory cap of 66,000",
    "•  H-2B statutory cap: 66,000 (set by Congress 1992, never raised by legislation). Supplemental cap: "
    "DHS authorized to add up to 64,716 additional visas per year via annual appropriations riders — "
    "both parties authorized this every year since FY2016",
    "•  Average certified H-2B job duration: 232 days (7.6 months). 53% of positions run 8+ months. "
    "28% run 9 months — the legal maximum. Program legally defined as temporary; operational reality "
    "is near-permanent.",
    "•  Top-10 H-2B occupations account for 98.8% of all approvals. Landscaping alone: 45.86%. Program "
    "concentrated in low-wage industries with documented wage theft histories.",
    "✓ SOURCE: EPI analysis of USCIS H-2B Employer Data Hub + DOL OFLC Performance Data FY2024 — "
    "epi.org/310379",
]

CHANGE_2 = [
    "•  H-2B certified wages below national OEWS average in ALL top-15 occupations FY2024 — the same "
    "dataset DOL uses to set H-2B wage minimums. Largest gaps: Forest workers 24.7% below / Meat "
    "poultry fish cutters 22.2% below / Construction laborers 15.4% below / Landscaping 10.7% below "
    "(37% of all H-2B jobs)",
    "•  Meatpacking H-2B certified wage: $16.66/hr vs. U.S.-born industry average $26.20/hr — H-2B "
    "workers certified at 63.6% of what U.S.-born workers earn in the same sector",
    "•  Employer private wage survey loophole: H-2B employers can substitute private surveys to set wages "
    "below local OEWS averages. Seafood industry documented use of this mechanism.",
    "✓ SOURCE: EPI analysis of BLS OEWS 2024 + DOL OFLC H-2B Performance Data FY2024 — epi.org/310379",
]

CHANGE_3 = [
    "•  H-2B industry wage theft FY2000–2024 (inflation-adjusted to 2024 dollars): $2,238,327,967 total "
    "back wages assessed across 7 major H-2B industries — $89.5 million stolen per year on average. "
    "1,822,164 workers assessed as wage theft victims. NOTE: Covers all workers in these industries "
    "(U.S.-born, green card holders, H-2B workers, unauthorized workers) — demonstrates wage "
    "suppression affects entire industry, not just visa holders.",
    "•  By industry: Construction $1.05B (avg $1,768/worker) / Food services $829M / Janitorial $119M / "
    "Hotels and motels $106M / Landscaping $77M / Amusement $40M / Forestry $14M",
    "•  WHD enforcement capacity: Only 611 investigators as of May 2025 policing 170 million workers — "
    "record low",
    "•  Meatpacking employer strategy (documented 1983–2024): As meatpacking unionization fell from "
    "37.4% (1983) to 13.5% (2024), foreign-born share rose to 42%+. Intentional employer strategy — "
    "aggressive recruitment of non-union immigrant labor to suppress bargaining power. Employers now "
    "lobbying Congress to expand H-2B to year-round meatpacking jobs — currently prohibited because "
    "H-2B requires temporary work.",
    "✓ SOURCE: DOL WHD Industries with High Prevalence of H-2B Workers FY2000–2024 — dol.gov/agencies/whd "
    "/ Workplace Justice Lab Rutgers May 2025 / EPI September 2025 — epi.org/310379",
]

CHANGE_4 = [
    "•  EPI September 2025 H-2B Report (Costa and Bivens) — epi.org/310379 — program size, wage "
    "suppression, $2.24B wage theft aggregate, meatpacking employer strategy",
    "•  DOL WHD Industries with High Prevalence of H-2B Workers FY2000–2024 — dol.gov/agencies/whd",
    "•  USCIS H-2B Employer Data Hub — uscis.gov/tools/reports-and-studies/h-2b-employer-data-hub",
    "•  Workplace Justice Lab Rutgers University May 2025 — 611 WHD investigators for 170M-worker "
    "labor market",
]


def main():
    print(f"Opening: {DOCX_PATH}")
    doc = Document(DOCX_PATH)
    print(f"  Paragraphs before: {len(doc.paragraphs)}")

    # Apply bottom-to-top to keep earlier indices stable
    insert_after(
        doc,
        "Centro de los Derechos del Migrante — cdmigrante.org/h2-visa-worker-rights/",
        CHANGE_4,
        "Change 4 (Section 07 verified sources)"
    )
    insert_after(
        doc,
        "2024 Farmworker Protection Rule suspended June 2025; proposed rescission July 2025",
        CHANGE_3,
        "Change 3 (Section 03b enforcement)"
    )
    insert_after(
        doc,
        "EPI: 2026 IFR wage cuts transfer $1.7",
        CHANGE_2,
        "Change 2 (Section 03b wages)"
    )
    insert_after(
        doc,
        "H-2B FY2026 Q1: ~39k certified; full-year routinely >200k",
        CHANGE_1,
        "Change 1 (Section 03b program scale)"
    )

    doc.save(DOCX_PATH)
    print(f"\n  Paragraphs after: {len(doc.paragraphs)}")
    print(f"  Saved: {DOCX_PATH}")


if __name__ == "__main__":
    main()
